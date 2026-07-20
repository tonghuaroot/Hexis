"""ReaderResult behavior: structured warnings, locator anchors, loud
truncation, fail-loud legacy formats, and end-to-end artifact preservation
through the pipeline."""

from __future__ import annotations

import hashlib
import os
import shutil
from pathlib import Path

import pytest

from services.ingest import Config, IngestionMode, IngestionPipeline, _hash_text
from services.ingest.readers import (
    DocxReader,
    PDFReader,
    ReaderResult,
    TextReader,
    XlsxReader,
    get_reader,
)
from tests.utils import _db_dsn, get_test_identifier

pytestmark = [pytest.mark.asyncio(loop_scope="session")]

_HAS_DOCX = True
try:
    import docx  # noqa: F401
except ImportError:
    _HAS_DOCX = False

_HAS_OPENPYXL = True
try:
    import openpyxl  # noqa: F401
except ImportError:
    _HAS_OPENPYXL = False

_HAS_PDF = True
try:
    import pdfplumber  # noqa: F401
except ImportError:
    _HAS_PDF = False


def test_default_read_result_wraps_legacy_read(tmp_path):
    p = tmp_path / "note.txt"
    p.write_text("plain text content")
    result = TextReader.read_result(p)
    assert isinstance(result, ReaderResult)
    assert result.text == "plain text content"
    assert result.warnings == []
    assert result.extractor_name == "TextReader"


def test_xls_fails_loud_with_conversion_hint(tmp_path):
    p = tmp_path / "legacy.xls"
    p.write_bytes(b"\xd0\xcf\x11\xe0 fake ole2")
    with pytest.raises(RuntimeError, match="convert it to .xlsx"):
        get_reader(p)


@pytest.mark.skipif(not _HAS_OPENPYXL, reason="openpyxl not installed")
def test_xlsx_truncation_warns_and_anchors_rows(tmp_path):
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Data"
    for i in range(10):
        ws.append([f"name{i}", i, i * 2.5])
    path = tmp_path / "book.xlsx"
    wb.save(path)

    old_cap = XlsxReader.MAX_ROWS_PER_SHEET
    XlsxReader.MAX_ROWS_PER_SHEET = 4
    try:
        result = XlsxReader.read_result(path)
    finally:
        XlsxReader.MAX_ROWS_PER_SHEET = old_cap

    truncations = [w for w in result.warnings if w.code == "truncated_rows"]
    assert len(truncations) == 1
    assert truncations[0].detail == {"sheet": "Data", "emitted": 4, "total": 10}
    assert "truncated: showing 4 of 10 rows" in result.text

    sheet_anchors = [a for a in result.anchors if a.kind == "sheet"]
    row_anchors = [a for a in result.anchors if a.kind == "row"]
    assert [a.value for a in sheet_anchors] == ["Data"]
    assert [a.value for a in row_anchors] == [1, 2, 3, 4]
    # Anchor offsets point at the exact line starts.
    for anchor in row_anchors:
        line_end = result.text.index("\n", anchor.char_offset)
        assert result.text[anchor.char_offset:line_end].startswith(f"name{anchor.value - 1}\t")


@pytest.mark.skipif(not _HAS_DOCX, reason="python-docx not installed")
def test_docx_heading_anchors(tmp_path):
    import docx

    document = docx.Document()
    document.add_heading("Alpha", level=1)
    document.add_paragraph("Alpha body text.")
    document.add_heading("Beta", level=2)
    document.add_paragraph("Beta body text.")
    path = tmp_path / "doc.docx"
    document.save(path)

    result = DocxReader.read_result(path)
    heading_anchors = [a for a in result.anchors if a.kind == "heading"]
    assert [a.value for a in heading_anchors] == [["Alpha"], ["Alpha", "Beta"]]
    for anchor in heading_anchors:
        assert result.text[anchor.char_offset:anchor.char_offset + len(anchor.value[-1])] == anchor.value[-1]
    assert result.extractor_name == "python-docx"


@pytest.mark.skipif(not _HAS_PDF, reason="pdfplumber not installed")
def test_pdf_page_anchors_and_ocr_warning_codes(tmp_path):
    # A minimal one-page text PDF written by hand (no reportlab dependency).
    pdf_bytes = (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]"
        b"/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
        b"4 0 obj<</Length 60>>stream\n"
        b"BT /F1 12 Tf 72 720 Td (Retention window is 90 days.) Tj ET\n"
        b"endstream endobj\n"
        b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
        b"trailer<</Root 1 0 R>>\n"
        b"%%EOF\n"
    )
    path = tmp_path / "doc.pdf"
    path.write_bytes(pdf_bytes)

    result = PDFReader.read_result(path)
    assert "[Page 1]" in result.text
    assert "Retention window is 90 days." in result.text
    page_anchors = [a for a in result.anchors if a.kind == "page"]
    assert [a.value for a in page_anchors] == [1]
    assert result.text[page_anchors[0].char_offset:].startswith("[Page 1]")
    # The page has a text layer, so no OCR/image warnings apply.
    assert not [w for w in result.warnings if w.code in ("ocr_used", "image_only_page")]


class _StubLLM:
    def __init__(self, marker: str):
        self.marker = marker
        self.call_count = 0

    async def complete_json(self, messages, temperature=0.2):
        self.call_count += 1
        text = str(messages[-1].get("content", ""))
        if "key 'items'" in text:
            return {"items": []}
        return {"valence": 0.0, "arousal": 0.2, "primary_emotion": "neutral",
                "intensity": 0.1, "summary": "Reference material."}


def _build_pipeline(marker: str) -> IngestionPipeline:
    config = Config(
        dsn=_db_dsn(os.environ.get("POSTGRES_DB")),
        llm_config={"provider": "openai", "model": "stub", "api_key": "stub"},
        mode=IngestionMode.FAST,
        verbose=False,
    )
    pipeline = IngestionPipeline(config)
    pipeline.llm = _StubLLM(marker)
    pipeline.appraiser.llm = pipeline.llm
    pipeline.extractor.llm = pipeline.llm
    return pipeline


async def test_ingest_file_preserves_artifact_and_records_run(db_pool, tmp_path):
    marker = get_test_identifier("artifactingest")
    path = tmp_path / f"{marker}.md"
    content = f"# Artifact Run {marker}\n\nOriginal bytes must be preserved before extraction."
    path.write_text(content)
    raw_sha = hashlib.sha256(path.read_bytes()).hexdigest()
    content_hash = _hash_text(content)

    pipeline = _build_pipeline(marker)
    try:
        await pipeline.ingest_file(path)
    finally:
        await pipeline.close()

    try:
        async with db_pool.acquire() as conn:
            row = await conn.fetchrow(
                """
                SELECT d.id AS doc_id, d.original_hash, a.sha256, a.storage_kind,
                       a.byte_size, a.original_filename,
                       (SELECT count(*) FROM source_extraction_runs r
                        WHERE r.source_document_id = d.id) AS run_count
                FROM source_documents d
                JOIN source_artifacts a ON a.source_document_id = d.id
                WHERE d.content_hash = $1
                """,
                content_hash,
            )
        assert row is not None
        assert row["sha256"] == raw_sha
        assert row["original_hash"] == raw_sha
        assert row["storage_kind"] == "database"
        assert row["original_filename"] == path.name
        assert row["run_count"] >= 1
    finally:
        async with db_pool.acquire() as conn:
            await conn.execute(
                "DELETE FROM memories WHERE source_attribution->>'content_hash' = $1", content_hash
            )
            await conn.execute("DELETE FROM source_artifacts WHERE sha256 = $1", raw_sha)
            await conn.execute("DELETE FROM source_documents WHERE content_hash = $1", content_hash)
            await conn.execute("DELETE FROM ingestion_receipts WHERE doc_ref = $1", content_hash)


async def test_failed_reader_still_preserves_artifact(db_pool, tmp_path):
    marker = get_test_identifier("failedreader")
    path = tmp_path / f"{marker}.docx"
    path.write_bytes(f"this is not a zip archive {marker}".encode())
    raw_sha = hashlib.sha256(path.read_bytes()).hexdigest()

    pipeline = _build_pipeline(marker)
    try:
        count = await pipeline.ingest_file(path)
    finally:
        await pipeline.close()

    try:
        assert count == 0
        async with db_pool.acquire() as conn:
            row = await conn.fetchrow(
                """
                SELECT a.id, a.source_document_id,
                       (SELECT r.status FROM source_extraction_runs r
                        WHERE r.artifact_id = a.id ORDER BY r.created_at DESC LIMIT 1) AS run_status
                FROM source_artifacts a WHERE a.sha256 = $1
                """,
                raw_sha,
            )
        assert row is not None, "artifact preserved even though extraction failed"
        assert row["source_document_id"] is None
        assert row["run_status"] == "failed"
    finally:
        async with db_pool.acquire() as conn:
            await conn.execute("DELETE FROM source_extraction_runs WHERE artifact_id = (SELECT id FROM source_artifacts WHERE sha256 = $1)", raw_sha)
            await conn.execute("DELETE FROM source_artifacts WHERE sha256 = $1", raw_sha)


async def test_large_artifact_goes_to_managed_directory(db_pool, tmp_path, monkeypatch):
    marker = get_test_identifier("bigartifact")
    artifact_dir = tmp_path / "artifact-store"
    monkeypatch.setenv("HEXIS_ARTIFACT_DIR", str(artifact_dir))

    path = tmp_path / f"{marker}.md"
    content = f"# Big {marker}\n\n" + ("filler paragraph text. " * 50)
    path.write_text(content)
    raw_sha = hashlib.sha256(path.read_bytes()).hexdigest()
    content_hash = _hash_text(content)

    pipeline = _build_pipeline(marker)
    pipeline.config.artifact_max_db_bytes = 16  # force the filesystem path
    try:
        await pipeline.ingest_file(path)
    finally:
        await pipeline.close()

    try:
        stored_file = artifact_dir / raw_sha[:2] / raw_sha
        assert stored_file.exists()
        assert stored_file.read_bytes() == path.read_bytes()
        async with db_pool.acquire() as conn:
            row = await conn.fetchrow(
                "SELECT storage_kind, storage_ref, bytes IS NULL AS no_db_bytes FROM source_artifacts WHERE sha256 = $1",
                raw_sha,
            )
        assert row["storage_kind"] == "filesystem"
        assert row["storage_ref"] == f"{raw_sha[:2]}/{raw_sha}"
        assert row["no_db_bytes"] is True
    finally:
        async with db_pool.acquire() as conn:
            await conn.execute(
                "DELETE FROM memories WHERE source_attribution->>'content_hash' = $1", content_hash
            )
            await conn.execute("DELETE FROM source_artifacts WHERE sha256 = $1", raw_sha)
            await conn.execute("DELETE FROM source_documents WHERE content_hash = $1", content_hash)
            await conn.execute("DELETE FROM ingestion_receipts WHERE doc_ref = $1", content_hash)
        shutil.rmtree(artifact_dir, ignore_errors=True)
